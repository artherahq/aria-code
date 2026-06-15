"""
file_analysis_tools.py — 多格式文件解析与内容提取层
=====================================================
支持格式：
  PDF          — pdfplumber (优先) 或 pypdf
  Word/DOCX    — python-docx
  Excel/XLSX   — openpyxl + pandas
  CSV/TSV      — pandas
  JSON/JSONL   — 内置 json
  Markdown/TXT — 直接读取
  图片          — PIL/Pillow 元数据 + base64 (发给视觉模型)
  HTML         — BeautifulSoup4 提取正文
  代码文件      — 语法感知提取 (py/js/ts/go/java/cpp 等)

全部函数返回统一格式：
  {"success": bool, "type": str, "content": str, "metadata": dict, ...}

依赖安装（可选，按需安装）：
  pip install pdfplumber python-docx openpyxl pandas pillow beautifulsoup4
"""

from __future__ import annotations

import base64
import json
import logging
import mimetypes
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ── Optional imports ──────────────────────────────────────────────────────────

def _try(mod):
    try:
        return __import__(mod)
    except ImportError:
        return None

# Detect available parsers at module load time (fast, no exceptions in hot path)
_pdfplumber = _try("pdfplumber")
_pypdf      = _try("pypdf")
_docx_mod   = _try("docx")
_openpyxl   = _try("openpyxl")
_pd         = _try("pandas")
_bs4        = _try("bs4")
_PIL        = _try("PIL")

# ── Data structures ───────────────────────────────────────────────────────────

@dataclass
class FileContent:
    """Normalised result of parsing any file."""
    success:   bool
    file_type: str          # "pdf" | "docx" | "xlsx" | "csv" | "json" | "image" | "code" | "text"
    path:      str
    filename:  str
    size_kb:   float
    content:   str          # Extracted text (may be truncated)
    metadata:  Dict[str, Any] = field(default_factory=dict)
    tables:    List[List[List[Any]]] = field(default_factory=list)  # nested: [sheet[row[cell]]]
    images_b64: List[str]   = field(default_factory=list)   # base64-encoded embedded images
    error:     Optional[str] = None
    truncated: bool = False
    char_count: int = 0


# ── Size / truncation limits ──────────────────────────────────────────────────

MAX_TEXT_CHARS  = 80_000   # ~20k tokens — safe for most context windows
MAX_TABLE_ROWS  = 200      # per sheet
MAX_PAGES       = 50       # PDF page limit
SUMMARY_CHARS   = 3_000    # for quick summary mode

# ── Main dispatcher ───────────────────────────────────────────────────────────

def parse_file(path_str: str, max_chars: int = MAX_TEXT_CHARS,
               include_images: bool = False) -> FileContent:
    """
    Parse any supported file and return normalised FileContent.

    path_str: absolute or ~-relative path
    max_chars: truncate text extraction at this many chars
    include_images: whether to base64-encode embedded images (expensive)
    """
    path = Path(path_str).expanduser().resolve()
    if not path.exists():
        return FileContent(False, "unknown", str(path), path.name, 0,
                           "", error=f"文件不存在: {path}")
    if not path.is_file():
        return FileContent(False, "unknown", str(path), path.name, 0,
                           "", error=f"不是文件: {path}")

    size_kb = path.stat().st_size / 1024
    if size_kb > 100 * 1024:  # 100 MB guard
        return FileContent(False, "unknown", str(path), path.name, size_kb,
                           "", error="文件过大（>100MB），请使用更小的文件")

    suffix = path.suffix.lstrip(".").lower()
    _DISPATCH = {
        "pdf":   _parse_pdf,
        "docx":  _parse_docx,
        "doc":   _parse_docx,
        "xlsx":  _parse_excel,
        "xls":   _parse_excel,
        "csv":   _parse_csv,
        "tsv":   _parse_csv,
        "json":  _parse_json,
        "jsonl": _parse_json,
        "html":  _parse_html,
        "htm":   _parse_html,
        "xml":   _parse_html,
        "md":    _parse_text,
        "txt":   _parse_text,
        "rst":   _parse_text,
        "log":   _parse_text,
        "yaml":  _parse_text,
        "yml":   _parse_text,
        "toml":  _parse_text,
        "ini":   _parse_text,
        "env":   _parse_text,
        "png":   _parse_image,
        "jpg":   _parse_image,
        "jpeg":  _parse_image,
        "gif":   _parse_image,
        "webp":  _parse_image,
        "bmp":   _parse_image,
    }
    # Code files
    _CODE_EXT = {"py","js","ts","tsx","jsx","go","java","c","cpp","h","hpp",
                 "rs","rb","php","swift","kt","scala","sh","bash","zsh",
                 "sql","r","m","cs","vb","lua","perl","ps1"}
    if suffix in _CODE_EXT:
        fn = _parse_code
    else:
        fn = _DISPATCH.get(suffix, _parse_text)

    try:
        result = fn(path, max_chars=max_chars, include_images=include_images)
        result.char_count = len(result.content)
        return result
    except Exception as e:
        logger.exception("parse_file failed for %s", path)
        return FileContent(False, suffix, str(path), path.name, size_kb,
                           "", error=f"解析失败: {e}")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    meta: Dict[str, Any] = {"pages": 0}
    text_parts = []
    tables = []

    if _pdfplumber:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            meta["pages"] = len(pdf.pages)
            meta["pdf_info"] = dict(pdf.metadata or {})
            for i, page in enumerate(pdf.pages[:MAX_PAGES]):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    text_parts.append(f"[第{i+1}页]\n{page_text}")
                # Tables
                for tbl in (page.extract_tables() or []):
                    if tbl:
                        tables.append(tbl)
    elif _pypdf:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        meta["pages"] = len(reader.pages)
        for i, page in enumerate(reader.pages[:MAX_PAGES]):
            t = page.extract_text() or ""
            if t.strip():
                text_parts.append(f"[第{i+1}页]\n{t.strip()}")
    else:
        return FileContent(False, "pdf", str(path), path.name, size_kb, "",
                           error="未安装 PDF 解析库，请运行: pip install pdfplumber")

    full_text = "\n\n".join(text_parts)
    truncated = len(full_text) > max_chars
    return FileContent(True, "pdf", str(path), path.name, size_kb,
                       full_text[:max_chars], meta, tables[:20],
                       truncated=truncated)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    if not _docx_mod:
        return FileContent(False, "docx", str(path), path.name, size_kb, "",
                           error="未安装 python-docx，请运行: pip install python-docx")
    import docx
    doc = docx.Document(str(path))

    parts = []
    tables = []

    # Core properties
    meta: Dict[str, Any] = {}
    try:
        cp = doc.core_properties
        meta = {
            "author":   cp.author,
            "created":  str(cp.created)[:10] if cp.created else "",
            "modified": str(cp.modified)[:10] if cp.modified else "",
            "title":    cp.title or "",
            "subject":  cp.subject or "",
        }
    except Exception:
        pass

    # Paragraphs
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        # Heading levels
        if para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "")
            parts.append(f"\n{'#' * int(level) if level.isdigit() else '##'} {t}")
        else:
            parts.append(t)

    # Tables
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows[:MAX_TABLE_ROWS]:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
            # Inline text representation
            header = " | ".join(rows[0])
            sep    = " | ".join(["---"] * len(rows[0]))
            body   = "\n".join(" | ".join(r) for r in rows[1:MAX_TABLE_ROWS])
            parts.append(f"\n[表格]\n{header}\n{sep}\n{body}")

    full_text = "\n".join(parts)
    truncated = len(full_text) > max_chars
    meta["paragraphs"] = len(doc.paragraphs)
    meta["tables"]     = len(doc.tables)
    return FileContent(True, "docx", str(path), path.name, size_kb,
                       full_text[:max_chars], meta, tables[:10],
                       truncated=truncated)


# ── Excel ─────────────────────────────────────────────────────────────────────

def _parse_excel(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    if not _pd:
        return FileContent(False, "xlsx", str(path), path.name, size_kb, "",
                           error="未安装 pandas，请运行: pip install pandas openpyxl")

    import pandas as pd
    try:
        sheets = pd.read_excel(str(path), sheet_name=None, engine="openpyxl",
                               nrows=MAX_TABLE_ROWS)
    except Exception as e:
        return FileContent(False, "xlsx", str(path), path.name, size_kb, "",
                           error=f"读取 Excel 失败: {e}")

    parts = []
    tables = []
    meta: Dict[str, Any] = {"sheets": list(sheets.keys()), "sheet_count": len(sheets)}

    for sname, df in sheets.items():
        df = df.fillna("").astype(str)
        rows = [list(df.columns)] + df.values.tolist()
        tables.append(rows)
        meta[f"sheet_{sname}_shape"] = f"{len(df)} 行 × {len(df.columns)} 列"

        # Text representation (first 50 rows shown)
        header = " | ".join(str(c) for c in df.columns)
        sep    = " | ".join(["---"] * len(df.columns))
        body_rows = df.head(50).values.tolist()
        body   = "\n".join(" | ".join(str(v)[:30] for v in r) for r in body_rows)
        parts.append(f"\n[Sheet: {sname}]  ({len(df)} 行 × {len(df.columns)} 列)\n"
                     f"{header}\n{sep}\n{body}")

        # Basic stats for numeric columns
        num_cols = df.select_dtypes(include="number") if hasattr(df, "select_dtypes") else None
        if num_cols is not None and not num_cols.empty:
            stats_lines = []
            for col in list(num_cols.columns)[:5]:
                s = num_cols[col].describe()
                stats_lines.append(f"  {col}: 均值={s.get('mean','')} 最大={s.get('max','')} 最小={s.get('min','')}")
            if stats_lines:
                parts.append("[数值统计]\n" + "\n".join(stats_lines))

    full_text = "\n".join(parts)
    truncated = len(full_text) > max_chars
    return FileContent(True, "xlsx", str(path), path.name, size_kb,
                       full_text[:max_chars], meta, tables,
                       truncated=truncated)


# ── CSV ───────────────────────────────────────────────────────────────────────

def _parse_csv(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    if not _pd:
        # Fallback: plain text
        return _parse_text(path, max_chars, include_images)

    import pandas as pd
    # Detect delimiter
    delim = "\t" if path.suffix.lower() == ".tsv" else ","
    try:
        df = pd.read_csv(str(path), sep=delim, nrows=MAX_TABLE_ROWS,
                         encoding="utf-8", on_bad_lines="skip")
    except Exception:
        try:
            df = pd.read_csv(str(path), sep=delim, nrows=MAX_TABLE_ROWS,
                             encoding="gbk", on_bad_lines="skip")
        except Exception as e:
            return FileContent(False, "csv", str(path), path.name, size_kb, "",
                               error=f"CSV 读取失败: {e}")

    meta: Dict[str, Any] = {
        "rows":    len(df),
        "columns": len(df.columns),
        "col_names": list(df.columns)[:30],
    }
    try:
        desc = df.describe(include="all").to_string()
        meta["stats_preview"] = desc[:1000]
    except Exception:
        pass

    rows = [list(df.columns)] + df.values.tolist()
    header = " | ".join(str(c) for c in df.columns)
    sep    = " | ".join(["---"] * len(df.columns))
    body   = "\n".join(" | ".join(str(v)[:25] for v in r) for r in df.head(80).values.tolist())
    full_text = f"[CSV: {path.name}]  {len(df)} 行 × {len(df.columns)} 列\n{header}\n{sep}\n{body}"
    truncated = len(full_text) > max_chars
    return FileContent(True, "csv", str(path), path.name, size_kb,
                       full_text[:max_chars], meta, [rows[:MAX_TABLE_ROWS]],
                       truncated=truncated)


# ── JSON ──────────────────────────────────────────────────────────────────────

def _parse_json(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    suffix = path.suffix.lower()

    try:
        if suffix == ".jsonl":
            lines = path.read_text(encoding="utf-8").splitlines()
            records = []
            for line in lines[:500]:
                line = line.strip()
                if line:
                    try: records.append(json.loads(line))
                    except Exception: pass
            data = records
            meta = {"format": "JSONL", "records": len(records), "sample": records[:3]}
        else:
            raw = path.read_text(encoding="utf-8")
            data = json.loads(raw)
            meta = {"format": "JSON", "type": type(data).__name__}
            if isinstance(data, list):
                meta["length"] = len(data)
                meta["sample"] = data[:3]
            elif isinstance(data, dict):
                meta["keys"] = list(data.keys())[:20]

        # Pretty-print (truncated)
        text = json.dumps(data, ensure_ascii=False, indent=2)
        truncated = len(text) > max_chars
        return FileContent(True, "json", str(path), path.name, size_kb,
                           text[:max_chars], meta, truncated=truncated)
    except Exception as e:
        return FileContent(False, "json", str(path), path.name, size_kb, "",
                           error=f"JSON 解析失败: {e}")


# ── HTML ──────────────────────────────────────────────────────────────────────

def _parse_html(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    raw = path.read_text(encoding="utf-8", errors="replace")
    meta: Dict[str, Any] = {}

    if _bs4:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw, "html.parser")
        # Remove scripts/styles
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        title = soup.find("title")
        meta["title"] = title.get_text().strip() if title else ""
        # Get main content
        main = soup.find("main") or soup.find("article") or soup.find("body") or soup
        text = main.get_text(separator="\n").strip()
        # Collapse blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
    else:
        # Simple tag stripping
        text = re.sub(r"<[^>]+>", " ", raw)
        text = re.sub(r"\s+", " ", text).strip()

    truncated = len(text) > max_chars
    return FileContent(True, "html", str(path), path.name, size_kb,
                       text[:max_chars], meta, truncated=truncated)


# ── Plain Text / Markdown / Code ──────────────────────────────────────────────

def _parse_text(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            text = path.read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = path.read_bytes().decode("utf-8", errors="replace")

    lines = text.count("\n")
    meta  = {"lines": lines, "encoding": "utf-8"}
    truncated = len(text) > max_chars
    return FileContent(True, "text", str(path), path.name, size_kb,
                       text[:max_chars], meta, truncated=truncated)


def _parse_code(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            text = path.read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = path.read_bytes().decode("utf-8", errors="replace")

    suffix = path.suffix.lstrip(".").lower()
    lines  = text.splitlines()
    meta: Dict[str, Any] = {
        "language":  suffix,
        "lines":     len(lines),
        "blank_lines": sum(1 for l in lines if not l.strip()),
    }

    # Extract function/class names
    _PATTERNS = {
        "py":   (r"^(?:async\s+)?def\s+(\w+)", r"^class\s+(\w+)"),
        "js":   (r"function\s+(\w+)\s*\(", r"class\s+(\w+)"),
        "ts":   (r"function\s+(\w+)\s*\(", r"class\s+(\w+)", r"interface\s+(\w+)"),
        "go":   (r"^func\s+\(?[^)]*\)?\s*(\w+)\s*\(", r"^type\s+(\w+)\s+struct"),
        "java": (r"(?:public|private|protected)?\s+\w+\s+(\w+)\s*\(", r"class\s+(\w+)"),
        "rs":   (r"^(?:pub\s+)?fn\s+(\w+)", r"^(?:pub\s+)?struct\s+(\w+)"),
    }
    patterns = _PATTERNS.get(suffix, [])
    symbols = []
    for pat in patterns:
        found = re.findall(pat, text, re.MULTILINE)
        symbols.extend(found[:20])
    if symbols:
        meta["symbols"] = list(dict.fromkeys(symbols))[:30]

    truncated = len(text) > max_chars
    return FileContent(True, "code", str(path), path.name, size_kb,
                       text[:max_chars], meta, truncated=truncated)


# ── Image ─────────────────────────────────────────────────────────────────────

def _parse_image(path: Path, max_chars: int, include_images: bool) -> FileContent:
    size_kb = path.stat().st_size / 1024
    suffix  = path.suffix.lstrip(".").lower()
    mime    = {"png":"image/png","jpg":"image/jpeg","jpeg":"image/jpeg",
               "gif":"image/gif","webp":"image/webp","bmp":"image/bmp"}.get(suffix,"image/png")
    meta: Dict[str, Any] = {"mime": mime, "size_kb": round(size_kb, 1)}

    if _PIL:
        from PIL import Image
        try:
            with Image.open(str(path)) as img:
                meta["width"]  = img.width
                meta["height"] = img.height
                meta["mode"]   = img.mode
                meta["format"] = img.format
                if hasattr(img, "_getexif") and img._getexif():
                    meta["has_exif"] = True
        except Exception:
            pass

    # Base64 for vision models
    images_b64 = []
    if include_images and size_kb < 10 * 1024:  # < 10MB
        try:
            b64 = base64.b64encode(path.read_bytes()).decode()
            images_b64 = [f"data:{mime};base64,{b64}"]
        except Exception:
            pass

    text = (f"[图片文件: {path.name}]\n"
            f"尺寸: {meta.get('width','?')}×{meta.get('height','?')} px\n"
            f"格式: {meta.get('format',suffix.upper())}\n"
            f"大小: {size_kb:.1f} KB")
    return FileContent(True, "image", str(path), path.name, size_kb,
                       text, meta, images_b64=images_b64)


# ── 多层分析提示词生成器 ──────────────────────────────────────────────────────

def build_analysis_prompt(fc: FileContent, layer: int = 1,
                           domain: str = "auto",
                           question: str = "") -> str:
    """
    为不同分析层次生成系统/用户提示词。

    layer 1 — 快速摘要    (300 字以内，Who/What/When/Why/Key metrics)
    layer 2 — 深度分析    (结构、要点、数据模式、异常值)
    layer 3 — 领域洞察    (财务/法律/技术/学术 — 由 domain 控制)
    layer 4 — 行动建议    (风险点、优化建议、下一步)
    """
    type_zh = {
        "pdf":   "PDF 文档",
        "docx":  "Word 文档",
        "xlsx":  "Excel 表格",
        "csv":   "CSV 数据文件",
        "json":  "JSON 数据",
        "image": "图片",
        "code":  "代码文件",
        "text":  "文本文件",
        "html":  "网页文件",
    }.get(fc.file_type, "文件")

    trunc_note = "\n\n⚠️ 注意：文件内容已被截断（超出上下文限制），以下为前段内容。" if fc.truncated else ""
    meta_summary = _format_meta(fc)

    # Auto-detect domain from content
    if domain == "auto":
        domain = _detect_domain(fc)

    _DOMAIN_CONTEXT = {
        "finance": "你是资深财务分析师，专注财报分析、现金流、盈利质量、风险敞口。",
        "legal":   "你是法律顾问，关注合同条款、风险条款、义务约束、免责声明。",
        "tech":    "你是高级软件工程师，评估代码质量、架构、安全、性能、可维护性。",
        "research":"你是研究员，提炼论文方法论、数据、结论、局限性及引用。",
        "realty":  "你是不动产分析师，关注物业数据、租金、估值、市场趋势、合规风险。",
        "medical": "你是医学顾问，总结诊断信息、治疗方案、用药风险（不构成诊断建议）。",
        "general": "你是专业文档分析师，全面理解文件内容。",
    }
    domain_ctx = _DOMAIN_CONTEXT.get(domain, _DOMAIN_CONTEXT["general"])

    layers = {
        1: f"""请对以下{type_zh}进行**快速摘要分析**（300字以内）：
{meta_summary}
要求：
1. 用一句话说明文件核心主题
2. 列出 3-5 个关键发现/数据点
3. 标注文件覆盖的时间范围（如有）
4. 指出最重要的结论或结果
{trunc_note}

---文件内容---
{fc.content[:SUMMARY_CHARS]}""",

        2: f"""请对以下{type_zh}进行**深度内容分析**：
{meta_summary}
分析维度：
1. **结构分析** — 文件章节/字段组织，逻辑流程
2. **数据要点** — 关键数字、趋势、比较（列表格式）
3. **异常与亮点** — 与常规预期显著不同的点
4. **信息完整性** — 是否有缺失、矛盾或模糊内容
5. **数据质量** — 若为数据文件：空值率、一致性
{trunc_note}

---文件内容---
{fc.content}""",

        3: f"""{domain_ctx}

请对以下{type_zh}进行**领域专项分析**：
{meta_summary}
重点关注：
1. **核心指标解读** — 本领域最重要的量化指标及其含义
2. **潜在风险** — 文件中反映或隐含的风险点
3. **与行业基准的偏差** — 什么是正常水平，当前数据如何？
4. **合规/规范性** — 是否符合本领域常规标准
5. **深层逻辑** — 表面数据背后的原因/驱动因素
{trunc_note}

---文件内容---
{fc.content}""",

        4: f"""{domain_ctx}

基于以下{type_zh}的内容，请给出**可执行的行动建议**：
{meta_summary}
要求：
1. **立即行动** — 需要立刻处理的事项（按优先级）
2. **改进建议** — 中期可以优化的方面
3. **风险预警** — 需要关注但未必立刻行动的隐患
4. **问题清单** — 文件不清晰/需要补充的 3-5 个问题
5. **下一步** — 建议的后续分析或决策步骤
{trunc_note}

---文件内容---
{fc.content[:max(len(fc.content)//2, SUMMARY_CHARS)]}""",
    }

    base_prompt = layers.get(layer, layers[1])
    if question:
        base_prompt = (f"关于以下{type_zh}，用户提问：\n\n**{question}**\n\n"
                       f"{meta_summary}{trunc_note}\n\n---文件内容---\n{fc.content}")
    return base_prompt


def _format_meta(fc: FileContent) -> str:
    """Format metadata for prompt header."""
    lines = [f"文件名: {fc.filename}", f"类型: {fc.file_type.upper()}",
             f"大小: {fc.size_kb:.1f} KB"]
    for k, v in fc.metadata.items():
        if k in ("pages","rows","columns","lines","language","sheets",
                 "records","length","sheet_count","paragraphs","tables",
                 "title","author","created"):
            lines.append(f"{k}: {v}")
    if fc.tables:
        lines.append(f"包含表格: {len(fc.tables)} 个")
    return "  ".join(lines[:8])


def _detect_domain(fc: FileContent) -> str:
    """Auto-detect analysis domain from content keywords."""
    text_lower = fc.content[:2000].lower()
    scores = {
        "finance":  sum(text_lower.count(k) for k in ["revenue","profit","ebitda","cashflow",
                        "净利润","营收","现金流","资产负债","毛利率","eps","roe","pe","财报"]),
        "legal":    sum(text_lower.count(k) for k in ["合同","甲方","乙方","违约","协议",
                        "条款","liability","indemnify","agreement","breach"]),
        "tech":     sum(text_lower.count(k) for k in ["def ","class ","function","import",
                        "select ","return ","var ","const ","type "]) + (2 if fc.file_type=="code" else 0),
        "research": sum(text_lower.count(k) for k in ["abstract","methodology","conclusion",
                        "hypothesis","摘要","结论","方法","样本量","显著性","p-value"]),
        "realty":   sum(text_lower.count(k) for k in ["租金","房价","物业","不动产","reit",
                        "产权","容积率","建筑面积","房地产","地块"]),
    }
    best = max(scores, key=scores.get)
    return best if scores[best] >= 2 else "general"


# ── Session file store (used by /file command) ────────────────────────────────

class FileSession:
    """
    Holds loaded files for the current REPL session.
    Provides multi-turn Q&A context injection.
    """
    def __init__(self):
        self._files: Dict[str, FileContent] = {}   # name → FileContent
        self._active: Optional[str] = None

    def load(self, path_str: str, include_images: bool = False) -> FileContent:
        fc = parse_file(path_str, include_images=include_images)
        if fc.success:
            self._files[fc.filename] = fc
            self._active = fc.filename
        return fc

    def get_active(self) -> Optional[FileContent]:
        if self._active and self._active in self._files:
            return self._files[self._active]
        return None

    def set_active(self, name: str) -> bool:
        if name in self._files:
            self._active = name
            return True
        # Partial match
        for key in self._files:
            if name.lower() in key.lower():
                self._active = key
                return True
        return False

    def list_files(self) -> List[Dict[str, Any]]:
        result = []
        for name, fc in self._files.items():
            result.append({
                "filename": name,
                "type":     fc.file_type,
                "size_kb":  round(fc.size_kb, 1),
                "chars":    fc.char_count,
                "active":   name == self._active,
                "truncated": fc.truncated,
            })
        return result

    def build_context_block(self, max_chars: int = 12_000) -> str:
        """Return a context block to inject into the system prompt."""
        fc = self.get_active()
        if not fc:
            return ""
        meta = _format_meta(fc)
        content_preview = fc.content[:max_chars]
        trunc = f"\n[内容已截断至前 {max_chars} 字符]" if fc.truncated or len(fc.content) > max_chars else ""
        return (f"\n\n---已加载文件: {fc.filename}---\n"
                f"{meta}\n\n{content_preview}{trunc}\n---文件结束---")

    def clear(self, name: Optional[str] = None):
        if name:
            self._files.pop(name, None)
            if self._active == name:
                self._active = next(iter(self._files), None)
        else:
            self._files.clear()
            self._active = None


# ── Dependency check ──────────────────────────────────────────────────────────

def check_parsers() -> Dict[str, bool]:
    return {
        "pdfplumber": _pdfplumber is not None,
        "pypdf":      _pypdf is not None,
        "python-docx": _docx_mod is not None,
        "pandas":     _pd is not None,
        "openpyxl":   _openpyxl is not None,
        "beautifulsoup4": _bs4 is not None,
        "Pillow":     _PIL is not None,
    }
